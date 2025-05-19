
import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Magazine PDF Feedback", layout="centered")
st.title("OCR NEA Magazine PDF Feedback Generator")

uploaded_file = st.file_uploader("Upload your magazine PDF", type=["pdf"])

def analyze_pdf(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text_content = ""
    image_count = 0
    for page in doc:
        text_content += page.get_text()
        image_count += len(page.get_images(full=True))
    word_count = len(text_content.split())
    excerpt = text_content[:500]
    return word_count, image_count, excerpt

def generate_feedback_doc(word_count, image_count):
    doc = Document()
    doc.add_heading('Magazine Submission Feedback', 0)
    doc.add_paragraph(f'Total word count in PDF: {word_count}')
    doc.add_paragraph(f'Total images detected: {image_count}')
    doc.add_paragraph('')

    if image_count >= 4:
        doc.add_paragraph("Good job! Your magazine contains the required minimum of 4 original images.")
    else:
        doc.add_paragraph("Warning: Your magazine should contain at least 4 original images as per the OCR NEA criteria.")

    if 250 <= word_count <= 350:
        doc.add_paragraph("The word count is within the expected range (around 300 words).")
    else:
        doc.add_paragraph("Note: The double-page spread should contain approximately 300 words. Consider adjusting the length.")

    doc.add_paragraph("Ensure the magazine has consistent use of color, typography, and layout for a strong house style.")
    doc.add_paragraph("Check the cover conventions such as masthead, cover lines, barcode, price, and edition details.")
    return doc

if uploaded_file is not None:
    st.success("PDF uploaded successfully!")
    word_count, image_count, excerpt = analyze_pdf(uploaded_file)
    st.write(f"**Estimated Word Count:** {word_count}")
    st.write(f"**Detected Images:** {image_count}")
    st.write("**Excerpt:**", excerpt[:500] + "...")

    feedback_doc = generate_feedback_doc(word_count, image_count)
    output = BytesIO()
    feedback_doc.save(output)
    output.seek(0)

    st.download_button(
        label="Download Feedback as Word Document",
        data=output,
        file_name="magazine_feedback.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
