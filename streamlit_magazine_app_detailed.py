
import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
import re

st.set_page_config(page_title="OCR Magazine Feedback", layout="centered")
st.title("📰 OCR NEA Magazine PDF Feedback Generator")

uploaded_file = st.file_uploader("Upload your magazine PDF", type=["pdf"])

def analyze_pdf(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text_content = ""
    image_count = 0
    page_count = len(doc)
    fonts_used = set()
    for page in doc:
        text_content += page.get_text()
        for block in page.get_text("dict")["blocks"]:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    fonts_used.add(span["font"])
    image_count = sum(len(page.get_images(full=True)) for page in doc)
    word_count = len(text_content.split())
    return word_count, image_count, text_content[:1000], page_count, sorted(fonts_used)

def detect_cover_conventions(text):
    checks = {
        "Masthead": bool(re.search(r"masthead", text, re.IGNORECASE)),
        "Cover lines": bool(re.search(r"cover line|headline|subheading", text, re.IGNORECASE)),
        "Barcode": bool(re.search(r"barcode", text, re.IGNORECASE)),
        "Price": bool(re.search(r"£|\$|\d+\.\d{2}", text)),
        "Edition info": bool(re.search(r"(January|February|Spring|Issue|Edition)", text, re.IGNORECASE)),
    }
    return checks

def generate_detailed_feedback(word_count, image_count, excerpt, page_count, fonts, conventions):
    doc = Document()
    doc.add_heading('Magazine Submission Feedback', 0)

    doc.add_paragraph(f"📝 Word count: {word_count}")
    doc.add_paragraph(f"🖼️ Number of images: {image_count}")
    doc.add_paragraph(f"📄 Page count: {page_count}")
    doc.add_paragraph(f"🔤 Fonts used: {', '.join(fonts)}")
    doc.add_paragraph("")

    if image_count >= 4:
        doc.add_paragraph("✅ Your magazine contains at least 4 original images.")
    else:
        doc.add_paragraph("⚠️ You need a minimum of 4 original images as per the OCR NEA Brief.")

    if 250 <= word_count <= 350:
        doc.add_paragraph("✅ The word count is within the expected range for a double-page spread.")
    else:
        doc.add_paragraph("⚠️ Adjust the content length — aim for around 300 words.")

    doc.add_paragraph("")
    doc.add_heading("🧠 House Style & Typography", level=1)
    if len(fonts) >= 2:
        doc.add_paragraph("✅ Multiple font styles detected — this supports a varied and effective house style.")
    else:
        doc.add_paragraph("⚠️ Only one font style detected — consider using at least 2 consistent fonts for contrast and hierarchy.")

    doc.add_paragraph("")
    doc.add_heading("📰 Cover Convention Check", level=1)
    for feature, present in conventions.items():
        symbol = "✅" if present else "❌"
        message = f"{symbol} {feature} detected." if present else f"{symbol} {feature} not clearly detected in text — check your layout visually."
        doc.add_paragraph(message)

    doc.add_paragraph("")
    doc.add_heading("📌 Excerpt from your PDF", level=1)
    doc.add_paragraph(excerpt[:800] + "..." if len(excerpt) > 800 else excerpt)

    return doc

if uploaded_file is not None:
    st.success("PDF uploaded successfully!")
    word_count, image_count, excerpt, page_count, fonts = analyze_pdf(uploaded_file)
    conventions = detect_cover_conventions(excerpt)

    st.write(f"**Word Count:** {word_count}")
    st.write(f"**Images Detected:** {image_count}")
    st.write(f"**Pages:** {page_count}")
    st.write(f"**Fonts Used:** {', '.join(fonts)}")

    st.write("### Cover Convention Summary")
    for k, v in conventions.items():
        st.write(f"- {k}: {'✅ Yes' if v else '❌ No'}")

    feedback_doc = generate_detailed_feedback(word_count, image_count, excerpt, page_count, fonts, conventions)
    output = BytesIO()
    feedback_doc.save(output)
    output.seek(0)

    st.download_button(
        label="📥 Download Detailed Feedback as Word Document",
        data=output,
        file_name="magazine_feedback_detailed.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
